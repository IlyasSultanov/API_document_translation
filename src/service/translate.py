import docx
import os
import tempfile
from fastapi import UploadFile, HTTPException



async def process_docx(file: UploadFile) -> str:
    """Извлекает текст из docx файла"""
    try:
        doc = docx.Document(file.file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {e}")


async def create_translated_docx(original_file: UploadFile, translated_text: str) -> str:
    """Создает новый docx файл с переведенным текстом, сохраняя форматирование"""
    try:
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, "translated.docx")

        doc = docx.Document(original_file.file)

        translated_paragraphs = translated_text.split("\n")
        for i, para in enumerate(doc.paragraphs):
            if i < len(translated_paragraphs):
                para.text = translated_paragraphs[i]

        doc.save(output_path)
        return output_path
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating translated DOCX: {e}")
