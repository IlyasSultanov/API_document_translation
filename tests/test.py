import pytest
from fastapi import UploadFile
from fastapi.testclient import TestClient
from main import app
from unittest.mock import patch, MagicMock
import docx
import os
import tempfile

client = TestClient(app)


@pytest.fixture
def sample_docx():
    """Фикстура для создания временного DOCX файла"""
    doc = docx.Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("This is a tests")
    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, "tests.docx")
    doc.save(file_path)
    yield file_path
    os.remove(file_path)


@pytest.fixture
def mock_translate():
    """Фикстура для мокирования API перевода"""
    with patch('src.connect_api.connect_api.requests.post') as mock_post:
        mock_response = MagicMock()
        mock_response.json.return_value = {"translated_text": "Привет мир\nЭто тест"}
        mock_response.raise_for_status.return_value = None
        mock_post.return_value = mock_response
        yield mock_post


def test_translate_docx_success(sample_docx, mock_translate):
    """Тест успешного перевода DOCX файла"""
    with open(sample_docx, "rb") as f:
        response = client.post(
            "/translate_docx/",
            files={"file": ("tests.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "translated_test.docx" in response.headers["content-disposition"]


def test_translate_docx_invalid_format():
    """Тест загрузки файла неверного формата"""
    response = client.post(
        "/translate_docx/",
        files={"file": ("tests.txt", b"tests content", "text/plain")}
    )
    assert response.status_code == 400
    assert "Only .docx files are accepted" in response.json()["detail"]


@pytest.mark.asyncio
async def test_process_docx(sample_docx):
    """Тест извлечения текста из DOCX"""
    from src.service.translate import process_docx
    with open(sample_docx, "rb") as f:
        text = await process_docx(UploadFile(filename="tests.docx", file=f))
    assert "Hello world" in text
    assert "This is a tests" in text


@pytest.mark.asyncio
async def test_translate_text():
    """Тест функции перевода текста"""
    from src.connect_api.connect_api import translate_text
    with patch('src.connect_api.connect_api.requests.post') as mock_post:
        mock_response = MagicMock()
        mock_response.json.return_value = {"translated_text": "Привет мир"}
        mock_post.return_value = mock_response

        result = await translate_text("Hello world")
        assert result == "Привет мир"


@pytest.mark.asyncio
async def test_create_translated_docx(sample_docx):
    """Тест создания переведенного DOCX"""
    from src.service.translate import create_translated_docx
    with open(sample_docx, "rb") as f:
        file = UploadFile(filename="tests.docx", file=f)
        output_path = await create_translated_docx(file, "Привет мир\nЭто тест")

        assert os.path.exists(output_path)
        doc = docx.Document(output_path)
        assert doc.paragraphs[0].text == "Привет мир"
        assert doc.paragraphs[1].text == "Это тест"
        os.remove(output_path)